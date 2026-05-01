"""Export the clean student version of the test as a .docx file.

Hides:
  - per-choice truth markers
  - skill labels (C1/C2/C3)

Renders LaTeX inline as plain text (KaTeX is for the Streamlit UI; Word
does not natively render KaTeX). Math segments delimited by `$ ... $`
are kept intact so a teacher can review them visually.
"""

from __future__ import annotations
from pathlib import Path
from typing import Iterable

from .models import Question

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:  # pragma: no cover
    Document = None  # type: ignore


def export_student_docx(
    questions: Iterable[Question],
    output_path: str | Path,
    title: str = "Géométrie dans l'espace — Test interactif",
) -> Path:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    output_path = Path(output_path)
    doc = Document()

    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.size = Pt(18)

    doc.add_paragraph(
        "Pour chaque question, cochez la (ou les) bonne(s) réponse(s). "
        "Vous pouvez répondre « Je ne sais pas »."
    )

    for q in questions:
        if q.is_diagnostic:
            doc.add_heading("Auto-évaluation", level=1)
        else:
            doc.add_heading(f"Question {q.number}", level=2)

        doc.add_paragraph(q.text)

        for i, choice in enumerate(q.choices):
            label = chr(ord("A") + i)
            doc.add_paragraph(f"☐ {label}. {choice.text}")

        if q.figure_path:
            note = doc.add_paragraph()
            note.add_run("(Figure associée à cette question.)").italic = True

        doc.add_paragraph()  # spacing

    doc.save(str(output_path))
    return output_path
