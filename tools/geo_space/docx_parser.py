"""Tolerant .docx parser for Geometry-in-Space MCQs.

The source document mixes paragraphs and tables; the format is fuzzy:
question codes can appear as "Q1", "Q 1", "Question 1"; truth markers
appear as "true / True / TRUE / vrai" anywhere in a choice line; skills
"C1", "C2", "C3" can appear inline. We strip everything decorative on
the way out so the resulting `Question` objects are clean.
"""

from __future__ import annotations
import re
from pathlib import Path
from typing import Iterable, List

from .models import Question, Choice
from .dataset import GROUP, LESSON

try:
    from docx import Document  # python-docx
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

# ─── Regexes ────────────────────────────────────────────────────────────────
RE_Q = re.compile(r"^\s*(?:Q(?:uestion)?\s*\.?\s*)(\d+)\b", re.IGNORECASE)
RE_SKILL = re.compile(r"\b(C[123])\b")
RE_TRUE = re.compile(r"\b(?:true|vrai)\b\.?", re.IGNORECASE)
RE_FALSE = re.compile(r"\b(?:false|faux)\b\.?", re.IGNORECASE)
RE_BULLET = re.compile(r"^\s*[\-\*\u2022\u25E6▪–—]+\s*")


def _iter_paragraphs(path: Path) -> Iterable[str]:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    doc = Document(str(path))
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        yield text


def clean_text(s: str) -> str:
    """Remove trailing markers (true/false, C1/C2/C3) and bullets."""
    s = RE_BULLET.sub("", s)
    s = RE_TRUE.sub("", s)
    s = RE_FALSE.sub("", s)
    s = RE_SKILL.sub("", s)
    s = re.sub(r"\s{2,}", " ", s).strip(" .;:,-")
    return s.strip()


def _is_truthy(s: str) -> bool:
    return bool(RE_TRUE.search(s))


def _detect_skill(s: str) -> str | None:
    m = RE_SKILL.search(s)
    return m.group(1) if m else None


def parse_docx(path: str | Path) -> List[Question]:
    """Parse questions from a Word document. Returns a list of `Question`.

    Recognised patterns:
    * "Q1" / "Q 1" / "Question 1" begin a new question.
    * Lines after the question stem are treated as choices until a new Q
      header or end of document.
    * "true" / "vrai" anywhere in a line marks the choice as correct.
    * "C1" / "C2" / "C3" anywhere in the question stem sets the skill.
    """
    path = Path(path)
    lines = list(_iter_paragraphs(path))

    questions: list[Question] = []
    current: Question | None = None

    for raw in lines:
        m = RE_Q.match(raw)
        if m:
            # close previous question
            if current is not None:
                _finalize(current)
                questions.append(current)
            n = int(m.group(1))
            stem_raw = raw[m.end():].strip(" .:-")
            skill = _detect_skill(stem_raw)
            current = Question(
                code=f"G{GROUP}-D{LESSON}-Q{n}",
                number=n,
                skill=skill,
                type_code=1,
                text=clean_text(stem_raw),
            )
            continue

        if current is None:
            continue

        # Continuation of stem or a new choice.
        if not current.choices and not RE_BULLET.match(raw) and len(raw) > 60:
            # Probably still part of the stem.
            current.text = clean_text(f"{current.text} {raw}")
            sk = _detect_skill(raw)
            if sk and not current.skill:
                current.skill = sk
            continue

        choice_text = clean_text(raw)
        if not choice_text:
            continue
        current.choices.append(Choice(text=choice_text, correct=_is_truthy(raw)))

    if current is not None:
        _finalize(current)
        questions.append(current)

    return questions


def _finalize(q: Question) -> None:
    q.is_multi = sum(c.correct for c in q.choices) > 1
    q.is_diagnostic = q.skill is None and not any(c.correct for c in q.choices)
