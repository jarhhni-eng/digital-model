# -*- coding: utf-8 -*-
"""
create_produit_scalaire_sample.py
==================================
Build a sample .docx matching the Produit Scalaire 1ere Bac document structure,
using the questions from the prompt.

Run:
    python create_produit_scalaire_sample.py
    python parse_produit_scalaire.py produit_scalaire_sample.docx --lesson 3
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_vanish(run):
    """Mark a run as hidden (w:vanish) so the marker is invisible when printed."""
    rPr = run._r.get_or_add_rPr()
    v = OxmlElement("w:vanish")
    rPr.append(v)


def add_q_table(doc, question_text: str,
                answers: list[tuple[str, str]],
                # answers: list of (display_text, marker)
                # marker: "true", "false", "skip", ""
                ):
    """
    Row 0  : question text
    Row 1+ : one answer per row
    """
    n = 1 + len(answers)
    tbl = doc.add_table(rows=n, cols=1)
    tbl.style = "Table Grid"

    tbl.rows[0].cells[0].text = question_text

    for i, (text, marker) in enumerate(answers):
        cell = tbl.rows[i + 1].cells[0]
        para = cell.paragraphs[0]
        para.add_run(text)
        if marker in ("true", "false", "correct", "faux"):
            r = para.add_run(f" {marker}")
            add_vanish(r)

    doc.add_paragraph("")


doc = Document()
doc.add_heading("Produit Scalaire - 1ere Bac Approfondissement", level=1)

# ── Type 1: Questions du cours ─────────────────────────────────────────────
doc.add_heading("Questions du cours", level=2)
doc.add_heading("Lecon 3", level=3)

# Q1 — self-assessment (no correct markers)
add_q_table(doc, "Q1 : a quel degre tu te rappelles la lecon du produit scalaire :",
    [
        ("J'ai tout oublie",        "skip"),
        ("Je me rappelle quelques parties", ""),
        ("Je me rappelle bien",     ""),
        ("Je me rappelle tout",     ""),
    ])

# Q2 C1
add_q_table(doc,
    "Q2 : Le produit scalaire des deux vecteurs U(a,c) et V(b,d) est le nombre reel note : C1",
    [
        ("||U||.||V||.cos((U,V))", "true"),
        ("||U||.||V||.sin((U,V))", "false"),
        ("ab+cd",                  "true"),
        ("J'ai tout oublie",       "skip"),
    ])

# Q3a C1
add_q_table(doc,
    "Q3 : Si l'un des vecteurs U ou V est nul, alors: C1",
    [
        ("U.V = 0",                 "true"),
        ("U.V = vecteur nul",       "false"),
        ("U et V sont orthogonaux", "false"),
        ("J'ai oublie",             "skip"),
    ])

# Q3b C1 (duplicate Q3 for perpendicular case)
add_q_table(doc,
    "Q3 : Si l'un des vecteurs U et V sont perpendiculaires alors: C1",
    [
        ("U.V = 0",           "true"),
        ("U.V = vecteur nul", "false"),
        ("U ou V est nul",    "false"),
        ("J'ai oublie",       "skip"),
    ])

# Q4 C3 (image: mediane.jpg)
add_q_table(doc,
    "Q4. Theoreme d'alkachy : C3\n"
    "BC^2 = AB^2 + AC^2 - 2(AB).(AC)",
    [
        ("BC^2 = AB^2 + AC^2 - 2(AB).(AC)", "true"),
        ("AB^2 = CB^2 + CA^2 - 2(CB).(CA)", "true"),
        ("J'ai oublie",                      "skip"),
    ])

# Q5 C3 (image: mediane.jpg)
add_q_table(doc,
    "Q5. Soit ABC un triangle, I est le milieu de BC, selon le Theoreme de la mediane: C3\n"
    "AB^2 + AC^2 = 1/2 BC^2 + 2 AI^2",
    [
        ("AB^2 + AC^2 = 1/2 BC^2 + 2 AI^2", "true"),
        ("AC^2 + BC^2 = 1/2 AB^2 + 2 BI^2", ""),
        ("J'ai oublie",                       "skip"),
    ])

# Q6 C1
add_q_table(doc,
    "Q6. det(U,V) sachant que U(a,b) et V(c,d) : C1",
    [
        ("Det = ac-bd", ""),
        ("Det = ab-bc", ""),
        ("Det = ad-bc", "true"),
        ("",            ""),
    ])

# Q7 (no skill)
add_q_table(doc,
    "Q7. Si le point M(x,y) appartient au cercle (C) de centre O(xo,yo) et de rayon 3cm :",
    [
        ("(x-xo)^2 + (y-yo)^2 = 3",    ""),
        ("(x-xo)^2 + (y-yo)^2 = 9",    "true"),
        ("x^2+y^2 = 2x+2y+9",           ""),
        ("J'ai oublie",                  "skip"),
    ])

# Q8 C1
add_q_table(doc,
    "Q8. Distance entre le point Omega et la Droite (D) : Omega(3;4) et (D): ax+by+c=0  C1",
    [
        ("D = |4a+3b+c| / sqrt(a^2+b^2)", ""),
        ("D = |3a+4b+c| / sqrt(a^2+c^2)", ""),
        ("D = |3a+4b+c| / sqrt(a^2+b^2)", "true"),
        ("J'ai oublie",                    "skip"),
    ])

# ── Type 2: Questions de la construction ──────────────────────────────────
doc.add_heading("Questions de la construction", level=2)

# Q9 C2 (image: repere.jpg)
add_q_table(doc,
    "Q9 : dans la fig 2 determiner les coordonnees des points B,C,D,E,F : C2",
    [
        ("B(4;4), C(3;0), D(9;3), E(10;6), F(12;5)", "true"),
        ("B(4;4), C(0;3), D(3;9), E(6;10), F(5;12)", ""),
        ("J'ai oublie",                                "skip"),
    ])

# Q10 C2 (image: repere.jpg)
add_q_table(doc,
    "Q10 : Determiner les coordonnees des vecteurs U, V, a, w: C2",
    [
        ("U(4;4), v(2;0), a(3;2), w(1;3)", "true"),
        ("U(4;4), v(0;2), a(2;3), w(3;1)", ""),
        ("J'ai oublie",                     "skip"),
    ])

# Q11 C2 (image: repere.jpg)
add_q_table(doc,
    "Q11 : le produit scalaire des vecteurs U.V (U=OB, V=OC) : C2",
    [
        ("U.V = ||OB||.||OC||.cos((OB;OC)) = 8",  "true"),
        ("U.V = ||OC||.||OB||.cos((OC;OB)) = -8", ""),
        ("U.V = ||OC||.||OB||.cos((OB;OB)) = 10", ""),
    ])

# Q12 C6 (image: repere.jpg)
add_q_table(doc,
    "Q12. Deduire la valeur de cos((OE;OC)) avec OE(10,6) et OC(3,0) : C6",
    [
        ("cos((DE;DF)) = 9/sqrt(34)",       "true"),
        ("cos((DE;DF)) = 9*sqrt(34)/34",    ""),
        ("cos((DE;DF)) = 3/sqrt(34)",       ""),
    ])

# Q13 C5 (image: repere.jpg)
add_q_table(doc,
    "Q13 : l'equation cartesienne de la droite (OB) est : C5",
    [
        ("x - y = 0",  ""),
        ("x + y = 0",  ""),
        ("x - 2y = 0", ""),
    ])

# Q14 C6 (image: repere.jpg)
add_q_table(doc,
    "Q14 : le produit scalaire des vecteurs a.w est : C6",
    [
        ("(DE).(DF) = 8",  "true"),
        ("(DE).(DF) = -8", ""),
        ("(DE).(DF) = 9",  ""),
    ])

# Q15 C5 (image: repere.jpg)
add_q_table(doc,
    "Q15 : le vecteur directeur de la droite (DE) a pour coordonnees : C5",
    [
        ("A(1;3)", ""),
        ("W(3;2)", ""),
        ("V(2;0)", ""),
    ])

# Q16 C5 (image: repere.jpg)
add_q_table(doc,
    "Q16 : tracer sur le repere orthonorme la droite (D): x+y-1=0  C5",
    [])   # open / drawing

# ── Type 3: Questions de raisonnement ─────────────────────────────────────
doc.add_heading("Questions de raisonnement", level=2)

# Q20 (was Q20 in original)
add_q_table(doc,
    "Q20 : la distance D entre le point A et la Droite ():",
    [
        ("15/sqrt(13)",       "true"),
        ("13/sqrt(15)",       ""),
        ("(15*sqrt(13))/13",  "true"),
        ("Je sais pas",       "skip"),
    ])

# Q21 C5
add_q_table(doc,
    "Q21 : l'equation de la droite (D) passant par A et perpendiculaire a () : C5",
    [
        ("x + 2y = 0",  ""),
        ("3x + 2y = 0", ""),
        ("Y = -3/2",    ""),
        ("Je sais pas", "skip"),
    ])

# Q22
add_q_table(doc,
    "Q22 : L'interpretation parametrique qui correspond a la droite () :",
    [
        ("{x=2+3t, y=3+2t}", ""),
        ("{x=2+3t, y=2+2t}", "true"),
        ("{x=3+3t, y=2+2t}", ""),
        ("Je sais pas",       "skip"),
    ])

# Q23
add_q_table(doc,
    "Q23 : les coordonnees de I, la projection orthogonale de A sur () :",
    [
        ("(6/13, -4/13)",   ""),
        ("(-4/13, -6/13)",  ""),
        ("(-4/13, 6/13)",   "true"),
        ("Je sais pas",      "skip"),
    ])

# Q24
add_q_table(doc,
    "Q24 : Determiner les coordonnees de B symetrique de A par rapport a () :",
    [
        ("Je sais pas", "skip"),
    ])

# Q25 C4
add_q_table(doc,
    "Q25. l'equation du cercle(C) de centre Omega(2;-1) et de rayon R=5 : C4",
    [
        ("(x-2)^2 + (y+1)^2 = 25",      "true"),
        ("x^2+y^2-4x+2y-20 = 0",        "true"),
        ("(x+2)^2 + (x-1)^2 = 25",      ""),
    ])

# Q26 C4
add_q_table(doc,
    "Q26. l'equation (x+2)^2 + (x-2)^2 = 16 est equivalente a : C4",
    [
        ("cercle de centre Omega(2;-2) et rayon R=4",   ""),
        ("cercle de centre Omega(-2;2) et rayon R=4",   "true"),
        ("cercle de centre Omega(2;-2) et rayon R=16",  ""),
    ])

# Q27 C3
add_q_table(doc,
    "(C) ensemble des points M(x;y), cercle de centre(1;2) et rayon=4 : C3",
    [
        ("x^2+y^2-2x+4y-16 = 0",  "false"),
        ("x^2+y^2-2x+2y-21 = 0",  "false"),
        ("x^2+y^2-2x+4y-21 = 0",  "correct"),
    ])

out = "produit_scalaire_sample.docx"
doc.save(out)
print(f"Sample created: {out}")
print("Now run:")
print(f"  python parse_produit_scalaire.py {out} --lesson 3")
