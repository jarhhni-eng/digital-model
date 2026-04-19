# -*- coding: utf-8 -*-
"""
parse_produit_scalaire.py
=========================
Parse a Word (.docx) document containing MCQ math questions for
1ère Bac — Produit Scalaire (Approfondissement).

Document structure expected
----------------------------
Type headers (above tables):
    "Questions du cours"          → Type 1
    "Questions de la construction" → Type 2
    "Questions de raisonnement"   → Type 3

Each question table:
    Row 0   : Question text  (may contain skill label Cₖ inline or hidden)
    Row 1+  : Answer options — each cell is one option.
              Correct marker  : "true", "vrai", "correct", "juste"
              Incorrect marker: "false", "faux"
              No marker       → incorrect
              Skip options    : "Je sais pas", "J'ai oublié", "J'ai tout oublié"

Skill labels C1–C6:
    C1. Parallélisme / perpendicularité via produit scalaire
    C2. Distances et angles via produit scalaire
    C3. Ensemble de points vérifiant MA·MB = 0
    C4. Centre et rayon d'un cercle via équation cartésienne
    C5. Représentation paramétrique / équation cartésienne
    C6. Propriétés analytiques du produit scalaire

Special image associations (hard-coded or auto-detected):
    Q4, Q5  → mediane.jpg
    Q9–Q16  → repere.jpg

Outputs
-------
    <name>_student.docx   — clean (no labels, no markers)
    <name>_teacher.docx   — full labels + correct answers + statistics

Usage
-----
    python parse_produit_scalaire.py document.docx
    python parse_produit_scalaire.py document.docx --lesson 3 --no-docx
"""

from __future__ import annotations

import argparse
import io
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# ── UTF-8 stdout on Windows ──────────────────────────────────────────────────
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required.  Install: pip install python-docx")
    sys.exit(1)

# ─── Constants ────────────────────────────────────────────────────────────────

TYPE_KEYWORDS: dict[int, list[str]] = {
    1: ["cours"],
    2: ["construction"],
    3: ["raisonnement"],
}

CORRECT_MARKERS: set[str] = {"true", "vrai", "correct", "juste"}
WRONG_MARKERS:   set[str] = {"false", "faux", "incorrect"}

# Options that always score 0 regardless of marking
SKIP_PHRASES: list[str] = [
    "je sais pas", "je ne sais pas",
    "j'ai oublié", "j ai oublié",
    "j'ai tout oublié", "j ai tout oublié",
]

SKILL_RE   = re.compile(r"\b(C[1-6])\b", re.IGNORECASE)
LESSON_RE  = re.compile(r"(?:le[cç]on|lesson|d)\s*(\d+)", re.IGNORECASE)
Q_CODE_RE  = re.compile(r"\b(Q\d+)\b", re.IGNORECASE)

# Images associated with specific Q-codes (original numbering in document)
IMAGE_MAP: dict[str, str] = {
    **{f"Q{n}": "repere.jpg"  for n in range(9, 17)},   # Q9–Q16
    "Q4": "mediane.jpg",
    "Q5": "mediane.jpg",
}

SKILL_DESCRIPTIONS: dict[str, str] = {
    "C1": "Parallélisme / perpendicularité via produit scalaire",
    "C2": "Distances et angles via produit scalaire",
    "C3": "Ensemble des points M vérifiant MA·MB = 0",
    "C4": "Centre et rayon d'un cercle via équation cartésienne",
    "C5": "Représentation paramétrique / équation cartésienne",
    "C6": "Propriétés analytiques du produit scalaire",
}

TYPE_LABELS: dict[int, str] = {
    1: "Type 1 - Questions du cours",
    2: "Type 2 - Questions de la construction",
    3: "Type 3 - Questions de raisonnement",
}

# ─── Data classes ─────────────────────────────────────────────────────────────

@dataclass
class AnswerOption:
    text: str
    is_correct: bool       # True = explicitly marked correct
    is_skip: bool          # True = "Je sais pas" / "J'ai oublié"

@dataclass
class Question:
    raw_id: str            # auto-numbered: T1-D3-Q2
    orig_code: str         # original Q-code in the document e.g. "Q4"
    type_num: int          # 1/2/3
    lesson_num: int
    q_num: int             # within (type, lesson)
    text: str
    skill: Optional[str]   # "C1"…"C6" or None
    options: list[AnswerOption] = field(default_factory=list)
    image: Optional[str]   = None     # filename e.g. "repere.jpg"
    is_open: bool          = False    # True if no selectable options

    @property
    def correct_indices(self) -> list[int]:
        return [i for i, o in enumerate(self.options) if o.is_correct]

    @property
    def max_score(self) -> float:
        if self.is_open:
            return float("nan")
        n = len(self.correct_indices)
        return float(n) if n else 1.0

# ─── Parsing helpers ──────────────────────────────────────────────────────────

def _all_text(elem) -> str:
    """All text inside an XML element, concatenated."""
    return "".join(e.text or "" for e in elem.iter(qn("w:t"))).strip()


def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _detect_skill(text: str) -> Optional[str]:
    m = SKILL_RE.search(text)
    return m.group(1).upper() if m else None


def _strip_skill(text: str) -> str:
    """Remove Cₖ labels from display text."""
    cleaned = SKILL_RE.sub("", text)
    return re.sub(r"\s{2,}", " ", cleaned).strip(" -:\u2013\u2014|")


def _strip_markers(text: str) -> str:
    """Remove correct/incorrect markers from display text."""
    pattern = r"\b(true|vrai|correct|juste|false|faux|incorrect)\b"
    cleaned = re.sub(pattern, "", text, flags=re.IGNORECASE)
    return re.sub(r"\s{2,}", " ", cleaned).strip(" -:\u2013|")


def _is_skip(text: str) -> bool:
    return any(ph in text.lower() for ph in SKIP_PHRASES)


def _cell_has_marker(cell, markers: set[str]) -> bool:
    """Check visible and vanish-hidden runs for marker words."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip().lower() in markers:
                return True
        for elem in para._element.iter(qn("w:t")):
            if elem.text and elem.text.strip().lower() in markers:
                return True
    return False


def _detect_type(text: str) -> Optional[int]:
    h = text.lower()
    for t, kws in TYPE_KEYWORDS.items():
        if any(kw in h for kw in kws):
            return t
    return None


def _detect_lesson(text: str) -> Optional[int]:
    m = LESSON_RE.search(text)
    return int(m.group(1)) if m else None


def _detect_q_code(text: str) -> Optional[str]:
    m = Q_CODE_RE.match(text.strip())
    return m.group(1).upper() if m else None


def _cell_has_image(cell) -> bool:
    return bool(cell._element.findall(".//" + qn("w:drawing")))

# ─── Main parser ──────────────────────────────────────────────────────────────

def parse_document(path: str, default_lesson: int = 1) -> list[Question]:
    doc = Document(path)
    questions: list[Question] = []

    current_type   = 1
    current_lesson = default_lesson
    # per (type, lesson) question counter
    q_counter: dict[tuple[int, int], int] = {}

    body_elems = list(doc.element.body)

    i = 0
    while i < len(body_elems):
        elem = body_elems[i]
        tag  = elem.tag.split("}")[-1]

        # ── Paragraph: look for type/lesson headings ──────────────────────
        if tag == "p":
            text = _all_text(elem)
            if text:
                t = _detect_type(text)
                if t is not None:
                    current_type = t
                l = _detect_lesson(text)
                if l is not None:
                    current_lesson = l
            i += 1
            continue

        # ── Table: one question ───────────────────────────────────────────
        if tag == "tbl":
            from docx.table import Table as DocxTable
            tbl  = DocxTable(elem, doc)
            rows = tbl.rows
            if not rows:
                i += 1
                continue

            # --- Row 0: question text ---
            q_text_raw = " ".join(
                _cell_text(c) for c in rows[0].cells
            ).strip()

            # Detect original Q-code (e.g. "Q4") from question text
            orig_code = _detect_q_code(q_text_raw) or ""

            # Skill
            q_skill = _detect_skill(q_text_raw)

            # Check for image in first row
            has_image = any(_cell_has_image(c) for c in rows[0].cells)

            # Determine associated image file
            image_file: Optional[str] = None
            if orig_code in IMAGE_MAP:
                image_file = IMAGE_MAP[orig_code]
            elif has_image:
                image_file = "figure.jpg"   # generic fallback

            # Clean question text
            q_text = _strip_markers(_strip_skill(q_text_raw))

            # Self-assessment / metacognitive questions (no meaningful MCQ)
            # Detect if question starts with "Comment estimez-vous" etc.
            is_metacog = any(
                kw in q_text_raw.lower()
                for kw in ["comment estimez-vous", "à quel degré"]
            )

            # --- Rows 1+: answer options ---
            options: list[AnswerOption] = []
            for row in rows[1:]:
                cells = row.cells
                if not cells:
                    continue
                # Merge cells in the row
                opt_raw  = " ".join(_cell_text(c) for c in cells).strip()
                if not opt_raw:
                    continue

                is_skip_opt = _is_skip(opt_raw)
                is_correct  = False
                for cell in cells:
                    if _cell_has_marker(cell, CORRECT_MARKERS):
                        is_correct = True
                        break

                # Clean display text
                opt_clean = _strip_markers(opt_raw)

                options.append(AnswerOption(
                    text=opt_clean,
                    is_correct=is_correct and not is_skip_opt,
                    is_skip=is_skip_opt,
                ))

            is_open = (len(options) == 0) or is_metacog or (
                all(not o.is_correct for o in options) and len(options) <= 1
            )

            # --- Build renumbered ID ---
            key = (current_type, current_lesson)
            q_counter[key] = q_counter.get(key, 0) + 1
            q_num  = q_counter[key]
            raw_id = f"T{current_type}-D{current_lesson}-Q{q_num}"

            questions.append(Question(
                raw_id=raw_id,
                orig_code=orig_code,
                type_num=current_type,
                lesson_num=current_lesson,
                q_num=q_num,
                text=q_text,
                skill=q_skill,
                options=options,
                image=image_file,
                is_open=is_open,
            ))

        i += 1

    return questions

# ─── Statistics ───────────────────────────────────────────────────────────────

@dataclass
class Stats:
    total_questions: int = 0
    total_max_score: float = 0.0
    by_type:   dict[int, dict] = field(default_factory=dict)
    by_lesson: dict[int, dict] = field(default_factory=dict)
    by_skill:  dict[str, dict] = field(default_factory=dict)


def compute_stats(questions: list[Question]) -> Stats:
    stats = Stats(total_questions=len(questions))

    for q in questions:
        ms = q.max_score
        is_nan = ms != ms

        if not is_nan:
            stats.total_max_score += ms

        # by type
        if q.type_num not in stats.by_type:
            stats.by_type[q.type_num] = {"count": 0, "max_score": 0.0, "skills": set()}
        stats.by_type[q.type_num]["count"] += 1
        if not is_nan:
            stats.by_type[q.type_num]["max_score"] += ms
        if q.skill:
            stats.by_type[q.type_num]["skills"].add(q.skill)

        # by lesson
        if q.lesson_num not in stats.by_lesson:
            stats.by_lesson[q.lesson_num] = {"count": 0, "max_score": 0.0, "types": set()}
        stats.by_lesson[q.lesson_num]["count"] += 1
        if not is_nan:
            stats.by_lesson[q.lesson_num]["max_score"] += ms
        stats.by_lesson[q.lesson_num]["types"].add(q.type_num)

        # by skill
        sk = q.skill or "Non defini"
        if sk not in stats.by_skill:
            stats.by_skill[sk] = {"count": 0, "max_score": 0.0, "questions": []}
        stats.by_skill[sk]["count"] += 1
        if not is_nan:
            stats.by_skill[sk]["max_score"] += ms
        stats.by_skill[sk]["questions"].append(q.raw_id)

    return stats

# ─── docx helpers ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_question_block(
    doc,
    q: Question,
    show_labels: bool,
    show_correct: bool,
):
    """Append one question block to a docx Document."""
    skill_txt  = f"  ({q.skill} - {SKILL_DESCRIPTIONS.get(q.skill, '')})" \
                 if show_labels and q.skill else ""
    header_txt = f"[{q.raw_id}]{skill_txt}  {q.text}"
    if show_labels and q.orig_code:
        header_txt = f"({q.orig_code}) {header_txt}"

    p   = doc.add_paragraph()
    run = p.add_run(header_txt)
    run.bold = True
    run.font.size = Pt(11)

    # Image reference
    if q.image:
        doc.add_paragraph(f"[Image : {q.image}]").italic = True

    if q.is_open:
        doc.add_paragraph("-> Question ouverte / dessin (correction manuelle)")
        doc.add_paragraph("")
        return

    # Options table
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"

    for idx, opt in enumerate(q.options):
        row = tbl.add_row()
        row.cells[0].text = chr(65 + idx) + "."
        row.cells[0].width = Inches(0.4)

        display = opt.text
        if show_correct and opt.is_correct:
            display += "  [CORRECT]"
        row.cells[1].text = display

        if show_correct and opt.is_correct:
            _set_cell_bg(row.cells[0], "D4EDDA")
            _set_cell_bg(row.cells[1], "D4EDDA")
        if opt.is_skip:
            _set_cell_bg(row.cells[0], "F0F0F0")
            _set_cell_bg(row.cells[1], "F0F0F0")

    doc.add_paragraph("")

# ─── Student version ──────────────────────────────────────────────────────────

def generate_student_version(questions: list[Question], out_path: str):
    doc = Document()
    doc.core_properties.title = "Version Eleve - Produit Scalaire (1ere Bac)"

    _add_heading(doc, "Evaluation - Version Eleve", level=1)
    _add_heading(doc, "Produit Scalaire - 1ere Bac Approfondissement", level=2)
    doc.add_paragraph(
        "Lisez attentivement chaque question. Pour les QCM, "
        "cochez la ou les reponses correctes."
    )
    doc.add_paragraph("")

    current_type   = None
    current_lesson = None

    for q in questions:
        if q.type_num != current_type:
            current_type = q.type_num
            _add_heading(doc, TYPE_LABELS[q.type_num], level=2)
        if q.lesson_num != current_lesson:
            current_lesson = q.lesson_num
            _add_heading(doc, f"Lecon {q.lesson_num}", level=3)

        _add_question_block(doc, q, show_labels=False, show_correct=False)

        # Page break between questions for "one per page" feel
        if not q.is_open:
            doc.add_paragraph("").add_run("").add_break()

    doc.save(out_path)
    print(f"  Student version: {out_path}")


# ─── Teacher version ──────────────────────────────────────────────────────────

def generate_teacher_version(questions: list[Question], stats: Stats, out_path: str):
    doc = Document()
    doc.core_properties.title = "Version Enseignant - Produit Scalaire (1ere Bac)"

    _add_heading(doc, "Evaluation - Version Enseignant / Analyse", level=1)
    _add_heading(doc, "Produit Scalaire - 1ere Bac Approfondissement", level=2)

    # ── Competency legend ──
    _add_heading(doc, "Competences ciblees", level=2)
    for code, desc in SKILL_DESCRIPTIONS.items():
        doc.add_paragraph(f"{code} : {desc}")
    doc.add_paragraph("")

    # ── Questions with answers ──
    current_type   = None
    current_lesson = None
    for q in questions:
        if q.type_num != current_type:
            current_type = q.type_num
            _add_heading(doc, TYPE_LABELS[q.type_num], level=2)
        if q.lesson_num != current_lesson:
            current_lesson = q.lesson_num
            _add_heading(doc, f"Lecon {q.lesson_num}", level=3)
        _add_question_block(doc, q, show_labels=True, show_correct=True)

    # ── Statistics ──
    doc.add_page_break()
    _add_heading(doc, "Statistiques de l'evaluation", level=1)

    doc.add_paragraph(f"Nombre total de questions : {stats.total_questions}")
    doc.add_paragraph(f"Score maximum total        : {stats.total_max_score:.1f} pt(s)")
    doc.add_paragraph("")

    # --- By type ---
    _add_heading(doc, "Par type de questions", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdrs = ["Type", "Nb questions", "Score max", "Competences"]
    for j, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1F3A8A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for t_num, data in sorted(stats.by_type.items()):
        row = tbl.add_row().cells
        row[0].text = TYPE_LABELS.get(t_num, f"Type {t_num}")
        row[1].text = str(data["count"])
        row[2].text = f"{data['max_score']:.1f}"
        row[3].text = ", ".join(sorted(data["skills"])) or "-"
    doc.add_paragraph("")

    # --- By lesson ---
    _add_heading(doc, "Par lecon", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    hdrs2 = ["Lecon", "Nb questions", "Score max", "Types"]
    for j, h in enumerate(hdrs2):
        cell = tbl2.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1F3A8A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for l_num, data in sorted(stats.by_lesson.items()):
        row = tbl2.add_row().cells
        row[0].text = f"Lecon {l_num}"
        row[1].text = str(data["count"])
        row[2].text = f"{data['max_score']:.1f}"
        row[3].text = ", ".join(f"T{t}" for t in sorted(data["types"]))
    doc.add_paragraph("")

    # --- By skill ---
    _add_heading(doc, "Par competence (Ck)", level=2)
    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = "Table Grid"
    hdrs3 = ["Competence", "Description", "Nb questions", "Score max", "Questions"]
    for j, h in enumerate(hdrs3):
        cell = tbl3.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1F3A8A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for sk, data in sorted(stats.by_skill.items()):
        row = tbl3.add_row().cells
        row[0].text = sk
        row[1].text = SKILL_DESCRIPTIONS.get(sk, "-")
        row[2].text = str(data["count"])
        row[3].text = f"{data['max_score']:.1f}"
        row[4].text = ", ".join(data["questions"])

    doc.save(out_path)
    print(f"  Teacher version: {out_path}")

# ─── Console report ───────────────────────────────────────────────────────────

def print_report(questions: list[Question], stats: Stats):
    sep = "-" * 72
    print(sep)
    print("RAPPORT D'ANALYSE - Produit Scalaire 1ere Bac")
    print(sep)
    print(f"Questions totales : {stats.total_questions}")
    print(f"Score max total   : {stats.total_max_score:.1f} pt(s)")
    print()

    for q in questions:
        cor = ", ".join(chr(65 + i) for i in q.correct_indices) \
              or ("(ouverte)" if q.is_open else "aucune")
        img = f"  img={q.image}" if q.image else ""
        print(
            f"  {q.raw_id:<18} ({q.orig_code or '  '})  "
            f"skill={q.skill or '-':<3}  "
            f"opts={len(q.options):<2}  correct={cor}{img}"
        )

    print()
    print("Par type :")
    for t, d in sorted(stats.by_type.items()):
        sk = ", ".join(sorted(d["skills"])) or "-"
        print(f"  T{t}: {d['count']} questions, max {d['max_score']:.1f} pt(s), skills: {sk}")

    print()
    print("Par lecon :")
    for l, d in sorted(stats.by_lesson.items()):
        print(f"  D{l}: {d['count']} questions, max {d['max_score']:.1f} pt(s)")

    print()
    print("Par competence :")
    for sk, d in sorted(stats.by_skill.items()):
        desc = SKILL_DESCRIPTIONS.get(sk, "")
        ids  = ", ".join(d["questions"])
        print(f"  {sk} ({desc[:40]})")
        print(f"       {d['count']} questions, max {d['max_score']:.1f} pt(s) - {ids}")

    print(sep)

# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Parse Produit Scalaire Word MCQ document → student + teacher docx."
    )
    parser.add_argument("docx_path", help="Path to source .docx file")
    parser.add_argument("--lesson",   "-l", type=int, default=1,
                        help="Default lesson number if not detected (default: 1)")
    parser.add_argument("--no-docx", action="store_true",
                        help="Print report only, skip docx generation")
    args = parser.parse_args()

    src = Path(args.docx_path)
    if not src.exists():
        print(f"ERROR: File not found: {src}")
        sys.exit(1)

    print(f"Parsing: {src}")
    questions = parse_document(str(src), default_lesson=args.lesson)

    if not questions:
        print("No questions found. Check document structure.")
        sys.exit(0)

    stats = compute_stats(questions)
    print_report(questions, stats)

    if not args.no_docx:
        stem    = src.stem
        out_dir = src.parent
        generate_student_version(questions, str(out_dir / f"{stem}_student.docx"))
        generate_teacher_version(questions, stats, str(out_dir / f"{stem}_teacher.docx"))


if __name__ == "__main__":
    main()
