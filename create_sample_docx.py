"""
create_sample_docx.py
=====================
Generates a sample .docx file matching the expected format for parse_word_questions.py,
using the "Droite dans le plan" (C5) example questions provided.

Run this to create a test document, then parse it:
  python create_sample_docx.py          # creates droite_sample.docx
  python parse_word_questions.py droite_sample.docx --lesson 5
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = Document()
doc.core_properties.title = "Droite dans le plan — Questions"

# ---------------------------------------------------------------------------
# Helper: add a hidden "true" marker in a run so the parser can detect it
# ---------------------------------------------------------------------------

def _add_run_with_hidden_marker(para, visible_text: str, marker: str):
    """
    Add visible_text as a normal run, then append a hidden run containing marker.
    The hidden run uses w:rPr/w:vanish so it is invisible when printed.
    """
    para.add_run(visible_text)
    hidden_run = para.add_run(f" {marker}")
    # Mark as hidden (vanish)
    rPr = hidden_run._r.get_or_add_rPr()
    vanish = rPr.get_or_add_vanish()  # type: ignore[attr-defined]


def _add_run_vanish(para, text: str):
    """Simpler helper: add a run and set vanish=True via raw XML."""
    from docx.oxml import OxmlElement
    run = para.add_run(text)
    rPr = run._r.get_or_add_rPr()
    vanish = OxmlElement("w:vanish")
    rPr.append(vanish)
    return run


def add_question_table(doc, question_text: str, answers: list[tuple[str, bool, bool]]):
    """
    answers: list of (text, is_correct, is_skip)
    Row 0: question
    Rows 1+: one row per answer
    """
    n_rows = 1 + len(answers)
    tbl = doc.add_table(rows=n_rows, cols=1)
    tbl.style = "Table Grid"

    # Row 0 — question
    tbl.rows[0].cells[0].text = question_text

    # Rows 1+ — answers
    for i, (text, is_correct, is_skip) in enumerate(answers):
        cell = tbl.rows[i + 1].cells[0]
        para = cell.paragraphs[0]
        para.add_run(text)
        if is_correct:
            _add_run_vanish(para, " true")

    doc.add_paragraph("")   # spacer


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

doc.add_heading("Leçon 5 — Droite dans le plan", level=1)

# ---- Type 1: Questions du cours ----
doc.add_heading("Questions du cours", level=2)

# Q1 — open / table task (C1)
q1_text = "C1  Complétez le tableau suivant pour la droite d'équation y = 2/3 x + 2/3"
add_question_table(doc, q1_text, [])   # open question, no options

# Q2 — drawing (C1)
q2_text = "C1  Tracez sur un repère orthonormal la droite dont vous avez complété le tableau."
add_question_table(doc, q2_text, [])   # drawing, no options

# Q3 — MCQ with two correct answers (C1)
q3_text = "C1  Calculez la distance D entre le point A(2;−3) et la droite (d) : 2x − 3y + 2 = 0"
add_question_table(doc, q3_text, [
    ("15/√13",         True,  False),
    ("13/√15",         False, False),
    ("(15√13)/13",     True,  False),
    ("Je sais pas",    False, True ),
])

# Q4 — MCQ, no "true" visible in sample (no correct marked)
q4_text = "Déterminez l'équation de la droite perpendiculaire à (d) passant par A(2;−3)"
add_question_table(doc, q4_text, [
    ("x + 2y = 0",   False, False),
    ("3x + 2y = 0",  False, False),
    ("Y = −3/2",     False, False),
    ("Je sais pas",  False, True ),
])

# Q5 — parametric interpretation (C1)
q5_text = "C1  Choisissez la représentation paramétrique correcte de la droite (d)"
add_question_table(doc, q5_text, [
    ("{x = 2 + 3t, y = 3 + 2t}",  False, False),
    ("{x = 2 + 3t, y = 2 + 2t}",  True,  False),
    ("{x = 3 + 3t, y = 2 + 2t}",  False, False),
    ("Je sais pas",                False, True ),
])

# ---- Type 2: Questions de la construction ----
doc.add_heading("Questions de la construction", level=2)

# Q6 — orthogonal projection (C2)
q6_text = "C2  Déterminez les coordonnées du projeté orthogonal H du point A(2;−3) sur (d)"
add_question_table(doc, q6_text, [
    ("(6/13, −4/13)",   False, False),
    ("(−4/13, −6/13)",  False, False),
    ("(−4/13, 6/13)",   True,  False),
    ("Je sais pas",     False, True ),
])

# Q7 — perpendicular line equation (C2)
q7_text = "C2  Écrivez l'équation réduite de la perpendiculaire à (d) passant par A(2;−3)"
add_question_table(doc, q7_text, [
    ("Y = −2/3 x + 2",  False, False),
    ("Y = 2/3 x + 2",   False, False),
    ("Y = −3/2 x + 2",  True,  False),
    ("Je sais pas",     False, True ),
])

out_path = "droite_sample.docx"
doc.save(out_path)
print(f"Sample document created: {out_path}")
print()
print("Now run:")
print(f"  python parse_word_questions.py {out_path} --lesson 5")
