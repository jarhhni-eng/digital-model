#!/usr/bin/env python3
"""
Builds questions_sample.docx that reproduces the exact content
from the user's specification so parser.py can be tested locally.

Run:
    pip install python-docx
    python build_sample_docx.py
    python parser.py questions_sample.docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    return p


def add_question(doc, q_text: str, options: list[tuple[str, str]]):
    """
    options: list of (option_text, label)
             label ∈ {"true","false","","J'ai oublié"}
    """
    table = doc.add_table(rows=3, cols=len(options))
    table.style = "Table Grid"

    # Row 0: merge all cells → question text
    row0 = table.rows[0]
    # merge manually by writing only to first cell (simple approach)
    row0.cells[0].text = q_text
    for i in range(1, len(options)):
        row0.cells[i].text = ""

    # Row 1: option texts
    row1 = table.rows[1]
    for i, (opt, _) in enumerate(options):
        row1.cells[i].text = opt

    # Row 2: labels (true/false/…)
    row2 = table.rows[2]
    for i, (_, lbl) in enumerate(options):
        row2.cells[i].text = lbl

    doc.add_paragraph()   # spacing


def build():
    doc = Document()

    # ── Lesson heading ────────────────────────────────────────────────────
    heading(doc, "1. Mémorisation : prérequis   1BAC SM", level=1)
    doc.add_paragraph()

    # ── Section: Questions du cours ───────────────────────────────────────
    heading(doc, "Questions du cours", level=2)
    doc.add_paragraph()

    # Q1 — no competency, no real right/wrong (survey item)
    add_question(doc,
        "Q1 : à quel degré tu te rappelles la leçon du produit scalaire ?",
        [
            ("J'ai tout oublié", ""),
            ("Je me rappelle quelques parties", ""),
            ("Je me rappelle bien", ""),
            ("Je me rappelle tout", ""),
        ]
    )

    # Q2 — C1
    add_question(doc,
        "Q2 : Le produit scalaire de U(a,c) et V(b,d) est : C1",
        [
            ("‖U⃗‖·‖V⃗‖ cos((U⃗,V⃗))", "true"),
            ("‖U⃗‖·‖V⃗‖ sin((U⃗,V⃗))", "false"),
            ("ab+cd", "true"),
            ("J'ai tout oublié", "J'ai oublié"),
        ]
    )

    # Q3a — C1
    add_question(doc,
        "Q3 : Si l'un des vecteurs U ou V est nul, alors : C1",
        [
            ("U⃗·V⃗ = 0", "true"),
            ("U⃗·V⃗ = 0⃗", "false"),
            ("U⃗ et V⃗ sont orthogonaux", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    # Q3b — C1 (duplicate code intentional)
    add_question(doc,
        "Q3 : Si U et V sont perpendiculaires, alors : C1",
        [
            ("U⃗·V⃗ = 0", "true"),
            ("U⃗·V⃗ = 0⃗", "false"),
            ("U⃗ ou V⃗ est nul", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    # Q4 — C3, image mediane.jpg
    add_question(doc,
        "Q4. Théorème d'Al-Kashi : C3\n"
        "Quels que soient les points A, B et C :",
        [
            ("BC² = AB² + AC² - 2·AB⃗·AC⃗", "true"),
            ("AB² = CB² + CA² - 2·CB⃗·CA⃗", "true"),
            ("J'ai oublié", "J'ai oublié"),
            ("", ""),
        ]
    )

    # Q5 — C3, image mediane.jpg
    add_question(doc,
        "Q5. Théorème de la médiane — I milieu de BC : C3",
        [
            ("AB² + AC² = ½ BC² + 2 AI²", "true"),
            ("AC² + BC² = ½ AB² + 2 BI²", "false"),
            ("J'ai oublié", "J'ai oublié"),
            ("", ""),
        ]
    )

    # Q6 — C1
    add_question(doc,
        "Q6. det(U,V) avec U(a,b) et V(c,d) : C1",
        [
            ("Det = ac-bd", "false"),
            ("Det = ab-bc", "false"),
            ("Det = ad-bc", "true"),
            ("", ""),
        ]
    )

    # Q7 — no competency
    add_question(doc,
        "Q7. Point M(x,y) appartient au cercle (C) de centre Ω(xₒ,yₒ) et rayon 3 :",
        [
            ("(x-xₒ)² + (y-yₒ)² = 3", "false"),
            ("(x-xₒ)² + (y-yₒ)² = 9", "true"),
            ("x² + y² = 2x + 2y + 9", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    # Q8 — C1
    add_question(doc,
        "Q8. Distance entre Ω(3;4) et (D) : ax+by+c=0 : C1",
        [
            ("d = |4a+3b+c| / √(a²+b²)", "false"),
            ("d = |3a+4b+c| / √(a²+c²)", "false"),
            ("d = |3a+4b+c| / √(a²+b²)", "true"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    # ── Section: Questions de la construction ─────────────────────────────
    heading(doc, "Questions de la construction", level=2)
    doc.add_paragraph()

    # Q9–Q16 would have image repere.jpg; add a couple as demo
    add_question(doc,
        "Q9. Dans le repère (O, I, J), le vecteur OA⃗ pour A(2,3) est : C2",
        [
            ("(2, 3)", "true"),
            ("(3, 2)", "false"),
            ("(−2, −3)", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    add_question(doc,
        "Q10. La distance entre A(1,2) et B(4,6) est : C2",
        [
            ("3", "false"),
            ("5", "true"),
            ("7", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    # ── Section: Questions de raisonnement ───────────────────────────────
    heading(doc, "Questions de raisonnement", level=2)
    doc.add_paragraph()

    add_question(doc,
        "Q11. Soit (D): 2x − y + 3 = 0.  La droite perpendiculaire passant par A(1,1) a pour équation : C1",
        [
            ("x + 2y − 3 = 0", "true"),
            ("2x − y + 1 = 0", "false"),
            ("x − 2y + 1 = 0", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    add_question(doc,
        "Q12. L'ensemble des points M vérifiant MA⃗·MB⃗ = 0 est : C3",
        [
            ("Un segment", "false"),
            ("Un cercle de diamètre [AB]", "true"),
            ("Une droite perpendiculaire à AB", "false"),
            ("J'ai oublié", "J'ai oublié"),
        ]
    )

    doc.save("questions_sample.docx")
    print("questions_sample.docx created — run: python parser.py questions_sample.docx")


if __name__ == "__main__":
    build()
