# -*- coding: utf-8 -*-
"""
parse_word_questions.py
=======================
Parse a Word (.docx) file containing MCQ math questions organized in tables.

Structure expected in the document
------------------------------------
- Section headers identify the question type:
    "Questions du cours"         → Type 1
    "Questions de la construction" → Type 2
    "Questions de raisonnement"  → Type 3
  (case-insensitive, partial match)

- Each question block is a table with:
    Row 0 : Question text (may contain a skill label such as "C1" or "C2"
            appearing as hidden/small text, or inline with the question)
    Rows 1+: Answer options — each row is one option.
             A cell that contains "true" or "faux"/"false" marks
             that option as correct.  "Je sais pas" is never correct.

- Lesson number is inferred from a heading such as
    "Leçon 1", "Leçon 2", … (or "D1", "D2", …) that appears before the tables.

Outputs
--------
Two docx files are produced next to the source file:
  <name>_student.docx   — clean: no labels, no correct markers
  <name>_teacher.docx   — labels + correct markers + statistics sheet

Usage
-----
  python parse_word_questions.py chemin/vers/document.docx
  python parse_word_questions.py chemin/vers/document.docx --lesson 5
"""

from __future__ import annotations

import argparse
import copy
import io
import os
import re
import sys

# Force UTF-8 on Windows terminals that default to cp1252
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is required.  Install with:  pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

TYPE_KEYWORDS: dict[int, list[str]] = {
    1: ["cours"],
    2: ["construction"],
    3: ["raisonnement"],
}

CORRECT_MARKERS: set[str] = {"true", "vrai", "correct", "juste"}
WRONG_MARKERS:   set[str] = {"false", "faux", "incorrect"}
SKIP_LABEL:      str       = "je sais pas"

SKILL_RE = re.compile(r"\b(C\d+)\b", re.IGNORECASE)
LESSON_RE = re.compile(r"(?:le[cç]on|lesson|d)\s*(\d+)", re.IGNORECASE)

# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class AnswerOption:
    text: str
    is_correct: bool          # True = marked correct in source
    is_skip: bool             # True = "Je sais pas"


@dataclass
class Question:
    raw_id: str               # e.g. "T1-D3-Q2"
    type_num: int             # 1 / 2 / 3
    lesson_num: int
    q_num: int                # within-lesson sequence
    text: str
    skill: Optional[str]      # "C1", "C2", … or None
    options: list[AnswerOption] = field(default_factory=list)
    has_image: bool = False   # True if the question references a figure
    is_open: bool = False     # True if no selectable options (drawing / table task)

    # ---------- scoring helpers ----------

    @property
    def correct_indices(self) -> list[int]:
        return [i for i, o in enumerate(self.options) if o.is_correct]

    def score(self, selected: list[int]) -> float:
        """
        For a given list of selected answer indices return the score.
        Rule: each correctly selected correct option → +1 pt,
              any wrong selection or missed correct selection → 0 for that slot.
        Open / drawing questions return NaN (manual grading).
        """
        if self.is_open or not self.options:
            return float("nan")
        correct = set(self.correct_indices)
        chosen  = set(selected)
        if not chosen:
            return 0.0
        # Partial scoring: credit only chosen options that are correct
        earned = len(correct & chosen)
        # Penalise wrong choices: reduce by wrong choices made
        wrong  = len(chosen - correct)
        return max(0.0, earned - wrong)

    @property
    def max_score(self) -> float:
        if self.is_open:
            return float("nan")
        return float(len(self.correct_indices)) if self.correct_indices else 1.0

# ---------------------------------------------------------------------------
# Parsing helpers
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    """Return all visible text from a table cell, stripped."""
    return " ".join(p.text for p in cell.paragraphs).strip()


def _detect_skill(text: str) -> Optional[str]:
    m = SKILL_RE.search(text)
    return m.group(1).upper() if m else None


def _strip_skill(text: str) -> str:
    """Remove skill labels (C1, C2 …) from text."""
    return SKILL_RE.sub("", text).strip(" -–—·:|")


def _cell_has_marker(cell, markers: set[str]) -> bool:
    """True if any run in the cell contains one of the markers (hidden or visible)."""
    for para in cell.paragraphs:
        for run in para.runs:
            token = run.text.strip().lower()
            if token in markers:
                return True
        # Also check raw XML for w:rPr hidden / vanish text
        for elem in para._element.iter(qn("w:t")):
            if elem.text and elem.text.strip().lower() in markers:
                return True
    return False


def _is_skip_option(text: str) -> bool:
    return SKIP_LABEL in text.lower()


def _cell_is_correct(cell) -> bool:
    return _cell_has_marker(cell, CORRECT_MARKERS)


def _cell_is_wrong(cell) -> bool:
    return _cell_has_marker(cell, WRONG_MARKERS)


def _detect_type(heading: str) -> Optional[int]:
    h = heading.lower()
    for t, kws in TYPE_KEYWORDS.items():
        if any(kw in h for kw in kws):
            return t
    return None


def _detect_lesson(heading: str) -> Optional[int]:
    m = LESSON_RE.search(heading)
    return int(m.group(1)) if m else None


def _paragraph_is_heading(para) -> bool:
    return para.style.name.lower().startswith("heading")


def _paragraph_text(para) -> str:
    return para.text.strip()

# ---------------------------------------------------------------------------
# Document parser
# ---------------------------------------------------------------------------

def parse_document(path: str, default_lesson: int = 1) -> list[Question]:
    """
    Walk through the document body and extract all questions.
    Returns a list of Question objects in document order.
    """
    doc = Document(path)

    questions: list[Question] = []
    current_type: int = 1
    current_lesson: int = default_lesson
    lesson_q_counter: dict[tuple[int, int], int] = {}  # (type, lesson) → count

    body_elements = list(doc.element.body)
    # We iterate over child elements of <w:body>; each is either a <w:p> or <w:tbl>
    i = 0
    while i < len(body_elements):
        elem = body_elements[i]
        tag  = elem.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            # Check paragraphs for type / lesson headings
            text = "".join(e.text or "" for e in elem.iter(qn("w:t"))).strip()
            if not text:
                i += 1
                continue
            t = _detect_type(text)
            if t is not None:
                current_type = t
            l = _detect_lesson(text)
            if l is not None:
                current_lesson = l

        elif tag == "tbl":
            # Each table = one question
            # Build a python-docx Table proxy to access cells easily
            from docx.table import Table as DocxTable
            tbl = DocxTable(elem, doc)

            rows = tbl.rows
            if not rows:
                i += 1
                continue

            # ---- Row 0: question text (all cells concatenated) ----
            q_text_parts = []
            q_skill: Optional[str] = None
            has_image = False

            for cell in rows[0].cells:
                ct = _cell_text(cell)
                if not q_skill:
                    q_skill = _detect_skill(ct)
                # Check for embedded images
                if cell._element.findall(".//" + qn("w:drawing")):
                    has_image = True
                q_text_parts.append(_strip_skill(ct))

            q_text = " ".join(p for p in q_text_parts if p).strip()

            # ---- Rows 1+: answer options ----
            options: list[AnswerOption] = []
            for row in rows[1:]:
                cells = row.cells
                if not cells:
                    continue
                # Merge all cells in the row for the option text
                opt_text = " ".join(_cell_text(c) for c in cells).strip()
                if not opt_text:
                    continue

                is_skip    = _is_skip_option(opt_text)
                is_correct = False
                for cell in cells:
                    if _cell_is_correct(cell):
                        is_correct = True
                        break

                # Strip markers from display text
                clean_opt = re.sub(
                    r"\b(true|vrai|correct|juste|false|faux|incorrect)\b",
                    "", opt_text, flags=re.IGNORECASE
                ).strip(" -–—·:|")

                options.append(AnswerOption(
                    text=clean_opt,
                    is_correct=is_correct,
                    is_skip=is_skip,
                ))

            is_open = len(options) == 0 or has_image

            # ---- Build question ID ----
            key = (current_type, current_lesson)
            lesson_q_counter[key] = lesson_q_counter.get(key, 0) + 1
            q_num   = lesson_q_counter[key]
            raw_id  = f"T{current_type}-D{current_lesson}-Q{q_num}"

            questions.append(Question(
                raw_id=raw_id,
                type_num=current_type,
                lesson_num=current_lesson,
                q_num=q_num,
                text=q_text,
                skill=q_skill,
                options=options,
                has_image=has_image,
                is_open=is_open,
            ))

        i += 1

    return questions

# ---------------------------------------------------------------------------
# Statistics
# ---------------------------------------------------------------------------

@dataclass
class Stats:
    total_questions: int = 0
    total_max_score: float = 0.0
    by_type: dict[int, dict]   = field(default_factory=dict)
    by_lesson: dict[int, dict] = field(default_factory=dict)
    by_skill: dict[str, dict]  = field(default_factory=dict)


def compute_stats(questions: list[Question]) -> Stats:
    stats = Stats(total_questions=len(questions))

    for q in questions:
        ms = q.max_score
        if not (ms != ms):  # not NaN
            stats.total_max_score += ms

        # by type
        if q.type_num not in stats.by_type:
            stats.by_type[q.type_num] = {"count": 0, "max_score": 0.0, "skills": set()}
        stats.by_type[q.type_num]["count"] += 1
        if not (ms != ms):
            stats.by_type[q.type_num]["max_score"] += ms
        if q.skill:
            stats.by_type[q.type_num]["skills"].add(q.skill)

        # by lesson
        if q.lesson_num not in stats.by_lesson:
            stats.by_lesson[q.lesson_num] = {"count": 0, "max_score": 0.0, "types": set()}
        stats.by_lesson[q.lesson_num]["count"] += 1
        if not (ms != ms):
            stats.by_lesson[q.lesson_num]["max_score"] += ms
        stats.by_lesson[q.lesson_num]["types"].add(q.type_num)

        # by skill
        sk = q.skill or "Non défini"
        if sk not in stats.by_skill:
            stats.by_skill[sk] = {"count": 0, "max_score": 0.0, "questions": []}
        stats.by_skill[sk]["count"] += 1
        if not (ms != ms):
            stats.by_skill[sk]["max_score"] += ms
        stats.by_skill[sk]["questions"].append(q.raw_id)

    return stats

# ---------------------------------------------------------------------------
# Document generators
# ---------------------------------------------------------------------------

TYPE_LABELS: dict[int, str] = {
    1: "Type 1 — Questions du cours",
    2: "Type 2 — Questions de la construction",
    3: "Type 3 — Questions de raisonnement",
}


def _add_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_question_to_doc(
    doc,
    q: Question,
    show_labels: bool,
    show_correct: bool,
):
    """Append one question block to a docx Document."""
    # Question header paragraph
    q_label = f"[{q.raw_id}]" if show_labels else f"[{q.raw_id}]"
    skill_txt = f"  ({q.skill})" if show_labels and q.skill else ""
    header_txt = f"{q_label}{skill_txt}  {q.text}"
    if q.has_image:
        header_txt += "  [voir figure]"

    p = doc.add_paragraph()
    run = p.add_run(header_txt)
    run.bold = True

    if q.is_open:
        doc.add_paragraph("-> Question ouverte / dessin (correction manuelle)")
        doc.add_paragraph("")
        return

    # Answer options table
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"

    for idx, opt in enumerate(q.options):
        row = tbl.add_row()
        # Column 0: letter label
        row.cells[0].text = chr(65 + idx) + "."
        row.cells[0].width = Inches(0.4)

        # Column 1: option text
        opt_display = opt.text
        if show_correct and opt.is_correct:
            opt_display += "  ✓"
        row.cells[1].text = opt_display

        if show_correct and opt.is_correct:
            _set_cell_bg(row.cells[0], "D4EDDA")   # light green
            _set_cell_bg(row.cells[1], "D4EDDA")
        elif show_correct and opt.is_skip:
            _set_cell_bg(row.cells[0], "F8F9FA")
            _set_cell_bg(row.cells[1], "F8F9FA")

    doc.add_paragraph("")   # spacing


def generate_student_version(questions: list[Question], out_path: str):
    """Clean document — no skill labels, no correct markers."""
    doc = Document()
    doc.core_properties.title = "Version Élève"

    _add_heading(doc, "Évaluation — Version Élève", level=1)
    doc.add_paragraph(
        "Lisez attentivement chaque question et choisissez la ou les réponses correctes."
    )

    current_type = None
    current_lesson = None

    for q in questions:
        if q.type_num != current_type:
            current_type = q.type_num
            _add_heading(doc, TYPE_LABELS[q.type_num], level=2)
        if q.lesson_num != current_lesson:
            current_lesson = q.lesson_num
            _add_heading(doc, f"Leçon {q.lesson_num}", level=3)

        _add_question_to_doc(doc, q, show_labels=False, show_correct=False)

    doc.save(out_path)
    print(f"  Student version: {out_path}")


def generate_teacher_version(questions: list[Question], stats: Stats, out_path: str):
    """Full document with labels, correct markers, and statistics."""
    doc = Document()
    doc.core_properties.title = "Version Enseignant / Analyse"

    _add_heading(doc, "Évaluation — Version Enseignant / Analyse", level=1)

    # ---- Questions with correct answers ----
    current_type   = None
    current_lesson = None

    for q in questions:
        if q.type_num != current_type:
            current_type = q.type_num
            _add_heading(doc, TYPE_LABELS[q.type_num], level=2)
        if q.lesson_num != current_lesson:
            current_lesson = q.lesson_num
            _add_heading(doc, f"Leçon {q.lesson_num}", level=3)

        _add_question_to_doc(doc, q, show_labels=True, show_correct=True)

    # ---- Statistics ----
    doc.add_page_break()
    _add_heading(doc, "Statistiques de l'évaluation", level=1)

    # Summary table
    doc.add_paragraph(f"Nombre total de questions : {stats.total_questions}")
    doc.add_paragraph(f"Score maximum total        : {stats.total_max_score:.1f} pt(s)")

    # By type
    _add_heading(doc, "Par type de questions", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Type"
    hdr[1].text = "Nb questions"
    hdr[2].text = "Score max"
    hdr[3].text = "Compétences ciblées"
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
        _set_cell_bg(h, "1F3A8A")
        h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for t_num, data in sorted(stats.by_type.items()):
        row = tbl.add_row().cells
        row[0].text = TYPE_LABELS.get(t_num, f"Type {t_num}")
        row[1].text = str(data["count"])
        row[2].text = f"{data['max_score']:.1f}"
        row[3].text = ", ".join(sorted(data["skills"])) or "—"

    doc.add_paragraph("")

    # By lesson
    _add_heading(doc, "Par leçon", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = "Leçon"
    hdr2[1].text = "Nb questions"
    hdr2[2].text = "Score max"
    hdr2[3].text = "Types inclus"
    for h in hdr2:
        h.paragraphs[0].runs[0].bold = True
        _set_cell_bg(h, "1F3A8A")
        h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for l_num, data in sorted(stats.by_lesson.items()):
        row = tbl2.add_row().cells
        row[0].text = f"Leçon {l_num}"
        row[1].text = str(data["count"])
        row[2].text = f"{data['max_score']:.1f}"
        row[3].text = ", ".join(f"T{t}" for t in sorted(data["types"]))

    doc.add_paragraph("")

    # By skill
    _add_heading(doc, "Par compétence (Cₖ)", level=2)
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    hdr3 = tbl3.rows[0].cells
    hdr3[0].text = "Compétence"
    hdr3[1].text = "Nb questions"
    hdr3[2].text = "Score max"
    hdr3[3].text = "Questions"
    for h in hdr3:
        h.paragraphs[0].runs[0].bold = True
        _set_cell_bg(h, "1F3A8A")
        h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for sk, data in sorted(stats.by_skill.items()):
        row = tbl3.add_row().cells
        row[0].text = sk
        row[1].text = str(data["count"])
        row[2].text = f"{data['max_score']:.1f}"
        row[3].text = ", ".join(data["questions"])

    doc.save(out_path)
    print(f"  Teacher version: {out_path}")

# ---------------------------------------------------------------------------
# Text report (stdout / plain file)
# ---------------------------------------------------------------------------

def print_report(questions: list[Question], stats: Stats):
    sep = "─" * 70
    print(sep)
    print("RAPPORT D'ANALYSE — QUESTIONS MCQ")
    print(sep)
    print(f"Total questions : {stats.total_questions}")
    print(f"Score max total : {stats.total_max_score:.1f} pt(s)")
    print()

    for q in questions:
        cor_labels = ", ".join(
            chr(65 + i) for i in q.correct_indices
        ) or ("(ouverte)" if q.is_open else "aucune")
        print(
            f"  {q.raw_id:<16}  skill={q.skill or '—':<5}  "
            f"options={len(q.options):<3}  correct={cor_labels}"
        )

    print()
    print("Par type :")
    for t, d in sorted(stats.by_type.items()):
        print(f"  T{t}: {d['count']} questions, max {d['max_score']:.1f} pt(s), "
              f"compétences: {', '.join(sorted(d['skills'])) or '—'}")

    print()
    print("Par leçon :")
    for l, d in sorted(stats.by_lesson.items()):
        print(f"  D{l}: {d['count']} questions, max {d['max_score']:.1f} pt(s)")

    print()
    print("Par compétence :")
    for sk, d in sorted(stats.by_skill.items()):
        print(f"  {sk}: {d['count']} questions, max {d['max_score']:.1f} pt(s) "
              f"— {', '.join(d['questions'])}")
    print(sep)

# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Parse a Word MCQ document and generate student / teacher versions."
    )
    parser.add_argument("docx_path", help="Path to the source .docx file")
    parser.add_argument(
        "--lesson", "-l",
        type=int, default=1,
        help="Default lesson number if not detected in the document (default: 1)"
    )
    parser.add_argument(
        "--no-docx", action="store_true",
        help="Skip generating .docx outputs (print report only)"
    )
    args = parser.parse_args()

    src = Path(args.docx_path)
    if not src.exists():
        print(f"ERROR: File not found — {src}")
        sys.exit(1)

    print(f"Parsing: {src}")
    questions = parse_document(str(src), default_lesson=args.lesson)

    if not questions:
        print("No questions found.  Check document structure (tables with headers).")
        sys.exit(0)

    stats = compute_stats(questions)
    print_report(questions, stats)

    if not args.no_docx:
        stem = src.stem
        out_dir = src.parent
        generate_student_version(questions, str(out_dir / f"{stem}_student.docx"))
        generate_teacher_version(questions, stats, str(out_dir / f"{stem}_teacher.docx"))


if __name__ == "__main__":
    main()
